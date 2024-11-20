import os

import pandas as pd
from docx import Document

def docx_to_csv(docx_path, csv_path):
    # Load the DOCX file
    doc = Document(docx_path)
    
    # Extract paragraphs
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip() != ""]
    
    # Create a DataFrame with each paragraph as a column
    df = pd.DataFrame(paragraphs[1:], columns=['Scenario'])
    print(f"{docx_path} :  size: {len(df)}")

    duplicated_rows = df[df.duplicated()]
    if len(duplicated_rows):
        print("Duplicated rows:")
        print(duplicated_rows)
    # Save DataFrame to a CSV file
    df.to_csv(csv_path, index=False)
    return df

def remove_matching_extensions_str(file_name, extensions):
    base_name = file_name
    for ext in extensions:
        if base_name.lower().endswith(ext.lower()):
            base_name = base_name[: -len(ext)]
    return base_name

# Function to remove the "_versions" suffix from the filename
def remove_versions_suffix(filename):
    return filename.replace('_versions', '')


def convert_files_to_csv(directory, extensions, output_directory):
    for root, dirs, files in os.walk(directory):
        for file_name in files:
            if file_name.endswith(extensions):
                file_path = os.path.join(root, file_name)
                print(f" File: {file_path}")

                docx_path = file_path
                os.makedirs(output_directory, exist_ok=True)

                # base_name= remove_matching_extensions_str(file_name, extensions)
                base_name, ext = os.path.splitext(file_name)
                csv_path = output_directory+remove_versions_suffix(base_name)+".csv"
                docx_to_csv(docx_path, csv_path)


def read_tables_from_docx(docx_path):
    # Load the DOCX file
    print(docx_path)
    doc = Document(docx_path)

    tables = []
    # Iterate over tables in the document
    for table in doc.tables:
        table_data = []
        # Iterate over rows in the table
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables

def convert_tables_to_dataframes(tables):
    dataframes = []
    for table in tables:
        df = pd.DataFrame(table)
        dataframes.append(df)
    return dataframes

def correction_in_column_names(path1):    
    files1 = os.listdir(path1)
    
    # Filter CSV files if needed
    csv_files1 = [os.path.join(path1,file) for file in files1 if file.endswith('.csv')]
    print(csv_files1)
    for file_path in csv_files1:
        df1 = pd.read_csv(file_path)
        if "Sub - mission" not in df1.columns:
            df1.insert(loc=2, column='Sub - mission', value="")
        df1['Sub - mission'] = ""
        rename_column_names = {'Hard Constraints': 'Hard Constrains' , 'Soft Constraints' : 'Soft Constrains (Preferences)'}
        df1.rename(columns=rename_column_names, inplace=True)

        # df1.drop(columns=['Unnamed: 0.1'], inplace=True)

        df1.to_csv(file_path,index= False)



def combine_scenarios_examples_and_version(df1,df2):
    columns_to_add = df1.columns[1:]
    values_to_add = df1.iloc[0, 1:]

    # Add the columns to df2
    for col in columns_to_add:
        df2[col] = values_to_add[col]
    return df2


def assign_labels_from_scenarios_examples(path1, path2, output_directory):

    os.makedirs(output_directory, exist_ok=True)

    files1 = os.listdir(path1)
    files2 = os.listdir(path2)
    # Filter CSV files if needed
    csv_files1 = [file for file in files1 if file.endswith('.csv')]
    
    csv_files2 = [file for file in files2 if file.endswith('.csv')]
    
    for file1 in csv_files1:
        for file2 in csv_files2:
            if file1 == file2:
                df1 = pd.read_csv(os.path.join(path1,file1))
                df2 = pd.read_csv(os.path.join(path2,file2))
                df = combine_scenarios_examples_and_version(df1,df2)
                df.to_csv(os.path.join(output_directory, file2), index = False)

    print(csv_files1)


def combine_scenarios(path):
    files = os.listdir(path)
    
    # Filter CSV files if needed
    csv_files = [os.path.join(path,file) for file in files if file.endswith('.csv')]
    dfs = [pd.read_csv(file_path) for file_path in csv_files]
    
    df = pd.concat(dfs,ignore_index=True)
    return df


def duplicates_in_each_scenarios(path):
    files = os.listdir(path)

    csv_files = [os.path.join(path,file) for file in files if file.endswith('.csv')]

    for file_path in csv_files:
        df = pd.read_csv(file_path)
        print(file_path)
        print(len(df), len(df.drop_duplicates()))



def combine_scenarios_from_different_sources(path1, path2, output_directory):
    
    os.makedirs(output_directory, exist_ok=True)

    files1 = os.listdir(path1)
    files2 = os.listdir(path2)


    csv_files1 = [file for file in files1 if file.endswith('.csv')]
    csv_files2 = [file for file in files2 if file.endswith('.csv')]

    for file_path1 in csv_files1:
        for file_path2 in csv_files2:
            
            if file_path1==file_path2:
                print(file_path1,file_path2)
                df1 = pd.read_csv(os.path.join(path1,file_path1))
                df2 = pd.read_csv(os.path.join(path2,file_path2))
                df = pd.concat([df1['Scenario'],df2['Scenario']], ignore_index=True)
                
               
                print(len(df), len(df.drop_duplicates()))

                df.drop_duplicates().to_csv(os.path.join(output_directory,file_path1),index = False)


def remove_duplicates_from_old_scenarios(path1 , path2):

    files1 = os.listdir(path1)
    files2 = os.listdir(path2)


    csv_files1 = [file for file in files1 if file.endswith('.csv')]
    csv_files2 = [file for file in files2 if file.endswith('.csv')]

    for file_path1 in csv_files1:
        for file_path2 in csv_files2:
            
            if file_path1==file_path2:
                print(file_path1,file_path2)

                df1 = pd.read_csv(os.path.join(path1,file_path1))
                df2 = pd.read_csv(os.path.join(path2,file_path2))

            
                df = pd.concat([df1['Scenario'],df2['Scenario']], ignore_index=True)
                print("Combined datapoint: ")
                print(f"Original: {len(df)}, Unique: {len(df.drop_duplicates())}")

                df = df2[~df2.Scenario.isin(df1.Scenario)]
                
                df.to_csv(os.path.join(path2,file_path2), index = False)
                
                print("Old datapoint")
                print(f"Original: {len(df)}, Unique: {len(df.drop_duplicates())}")

                print()


def train_test_split_without_suffle(path1, path2, path3, train = 28):
    os.makedirs(path2, exist_ok=True)
    os.makedirs(path3, exist_ok=True)

    files1 = os.listdir(path1)
    csv_files1 = [file for file in files1 if file.endswith('.csv')]

    for file_path1 in csv_files1:
        df1 = pd.read_csv(os.path.join(path1,file_path1))
        df1[:train].to_csv(os.path.join(path2,file_path1), index= False)
        df1[train:].to_csv(os.path.join(path3,file_path1), index= False)



def remove_label(path1, path2):
    os.makedirs(path2, exist_ok=True)

    files1 = os.listdir(path1)
    csv_files1 = [file for file in files1 if file.endswith('.csv')]

    for file_path1 in csv_files1:

        df1 = pd.read_csv(os.path.join(path1,file_path1))
        df1["Scenario"].to_csv(os.path.join(path2,file_path1), index= False)


def list_files_and_folders(directory):
    for root, dirs, files in os.walk(directory):
        print(f"Root: {root}")
        for dir_name in dirs:
            print(f" Directory: {dir_name}")
        for file_name in files:
            print(f" File: {file_name}")